import os
from datetime import datetime
from ..models.domain import DiligenceReport, MarketData, CompetitorLandscape, FinancialProfile, LegalProfile

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from jinja2 import Template
    HAS_JINJA = True
except ImportError:
    HAS_JINJA = False

class ReportRenderer:
    """Class responsible for generating Markdown, DOCX, and HTML reports based on DiligenceReport domain data."""

    def __init__(self, output_dir: str):
        """Initialize renderer and ensure output directory exists."""
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def render_markdown(self, report: DiligenceReport) -> str:
        """Generate full Markdown report string."""
        badge = self._render_score_badge(report.investment_score)
        formatted_date = datetime.now().strftime("%B %d, %Y")
        
        toc = (
            "## Table of Contents\n"
            "1. [Executive Summary](#executive-summary)\n"
            "2. [Risk Assessment](#risk-assessment)\n"
            "3. [Market Opportunity](#market-opportunity)\n"
            "4. [Competitive Landscape](#competitive-landscape)\n"
            "5. [Financial Profile](#financial-profile)\n"
            "6. [Legal Standing](#legal-standing)\n"
            "7. [Sources & Bibliography](#sources-bibliography)\n"
            "8. [Diligence Pipeline Metadata](#diligence-pipeline-metadata)\n"
        )
        
        metadata_rows = []
        for r in report.agent_results:
            metadata_rows.append(f"| {r.agent_name} | {r.status} | {r.duration_ms} ms | {r.error or 'None'} |")
            
        metadata_table = (
            "| Agent Name | Status | Execution Time | Error Info |\n"
            "| --- | --- | --- | --- |\n" +
            "\n".join(metadata_rows)
        )
        
        # Gather sources
        all_sources = []
        for r in report.agent_results:
            if r.sources:
                all_sources.extend(r.sources)
        sources_list = "\n".join(f"- {src}" for src in sorted(list(set(all_sources))) if src.startswith("http"))
        if not sources_list:
            sources_list = "No external data source URLs recorded."
        
        return (
            f"# VentureMind Due Diligence Report: {report.startup_name}\n\n"
            f"**Generated On:** {formatted_date}\n"
            f"**Investment Recommendation:** {badge}\n\n"
            f"{toc}\n\n"
            "### Executive Summary\n\n"
            f"{report.summary}\n\n"
            "### Risk Assessment\n\n"
            f"{self._format_risk_flags(report.risk_flags)}\n\n"
            f"{self._render_market_section(report.market)}\n\n"
            f"{self._render_competitor_section(report.competitors)}\n\n"
            f"{self._render_financial_section(report.financials)}\n\n"
            f"{self._render_legal_section(report.legal)}\n\n"
            "### Sources & Bibliography\n\n"
            f"{sources_list}\n\n"
            "### Diligence Pipeline Metadata\n\n"
            f"{metadata_table}\n"
        )

    def render_docx(self, report: DiligenceReport) -> str:
        """Generate DOCX file and return its filepath."""
        if not HAS_DOCX:
            raise ImportError(
                "The python-docx library is not installed. "
                "Please run `pip install python-docx` to enable DOCX report generation."
            )
            
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = report.startup_name.lower().replace(' ', '_')
        filepath = os.path.join(self.output_dir, f"{safe_name}_{timestamp}.docx")
        
        doc = docx.Document()
        doc.add_heading(f"VentureMind Due Diligence Report: {report.startup_name}", 0)
        
        badge = self._render_score_badge(report.investment_score)
        doc.add_paragraph(f"Generated On: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Recommendation: {badge}")
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.summary)
        
        # Risk Assessment
        doc.add_heading("Risk Assessment", level=1)
        if not report.risk_flags:
            doc.add_paragraph("No significant risk flags identified.")
        else:
            for flag in report.risk_flags:
                doc.add_paragraph(f"• {flag}")
            
        # Market
        if report.market:
            doc.add_heading("Market Opportunity", level=1)
            doc.add_paragraph(f"CAGR: {report.market.cagr_pct:.1f}%")
            
            doc.add_heading("Market Sizing", level=2)
            table = doc.add_table(rows=4, cols=2)
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "Value (USD)"
            table.cell(1, 0).text = "TAM"
            table.cell(1, 1).text = f"${report.market.tam_usd:,.2f}"
            table.cell(2, 0).text = "SAM"
            table.cell(2, 1).text = f"${report.market.sam_usd:,.2f}"
            table.cell(3, 0).text = "SOM"
            table.cell(3, 1).text = f"${report.market.som_usd:,.2f}"
            
            doc.add_heading("Key Industry Trends", level=2)
            for trend in report.market.key_trends:
                doc.add_paragraph(f"• {trend}")
        else:
            doc.add_heading("Market Opportunity", level=1)
            doc.add_paragraph("Data unavailable.")
                
        # Competitor
        if report.competitors:
            doc.add_heading("Competitive Landscape", level=1)
            differentiation_pct = int(report.competitors.differentiation_score * 100)
            doc.add_paragraph(f"Differentiation Score: {differentiation_pct}%")
            doc.add_paragraph(f"Positioning Summary: {report.competitors.positioning_summary}")
            
            doc.add_heading("Direct Competitor Matrix", level=2)
            if not report.competitors.competitors:
                doc.add_paragraph("No direct competitors identified.")
            else:
                table = doc.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Competitor"
                hdr_cells[1].text = "Founded"
                hdr_cells[2].text = "Funding"
                hdr_cells[3].text = "Market Share"
                
                for c in report.competitors.competitors:
                    row_cells = table.add_row().cells
                    row_cells[0].text = c.name
                    row_cells[1].text = str(c.founded_year)
                    row_cells[2].text = f"${c.funding_usd:,.2f}" if c.funding_usd else "unknown"
                    row_cells[3].text = f"{c.market_share_pct:.1f}%" if c.market_share_pct else "unknown"
        else:
            doc.add_heading("Competitive Landscape", level=1)
            doc.add_paragraph("Data unavailable.")
                
        # Financials
        if report.financials:
            doc.add_heading("Financial Profile", level=1)
            doc.add_paragraph(f"Estimated Burn: {report.financials.burn_rate_estimate}")
            doc.add_paragraph(f"Estimated Runway: {report.financials.runway_estimate}")
            
            doc.add_heading("Financial Signals & Metrics", level=2)
            if not report.financials.signals:
                doc.add_paragraph("No financial signals recorded.")
            else:
                table = doc.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Metric"
                hdr_cells[1].text = "Value"
                hdr_cells[2].text = "Period"
                hdr_cells[3].text = "Source"
                hdr_cells[4].text = "Confidence"
                
                for s in report.financials.signals:
                    row_cells = table.add_row().cells
                    row_cells[0].text = s.metric_name
                    row_cells[1].text = s.value
                    row_cells[2].text = s.period
                    row_cells[3].text = s.source
                    row_cells[4].text = f"{s.confidence:.2f}"
            
            doc.add_heading("Funding Rounds", level=2)
            if not report.financials.funding_rounds:
                doc.add_paragraph("No funding rounds recorded.")
            else:
                for r in report.financials.funding_rounds:
                    doc.add_paragraph(f"• {r.get('round')}: {r.get('amount_str')} ({r.get('date_str')})")
        else:
            doc.add_heading("Financial Profile", level=1)
            doc.add_paragraph("Data unavailable.")
                
        # Legal
        if report.legal:
            doc.add_heading("Legal Standing", level=1)
            doc.add_paragraph(f"Incorporation: {report.legal.incorporation_status}")
            doc.add_paragraph(f"Patents: {report.legal.patent_count}")
            doc.add_paragraph(f"Trademarks: {report.legal.trademark_count}")
            
            doc.add_heading("Legal & Compliance Flags", level=2)
            if not report.legal.flags:
                doc.add_paragraph("No active legal or regulatory flags identified.")
            else:
                for f in report.legal.flags:
                    doc.add_paragraph(f"• [{f.flag_type} - {f.severity}] {f.description}")
        else:
            doc.add_heading("Legal Standing", level=1)
            doc.add_paragraph("Data unavailable.")
            
        # Sources & Bibliography
        doc.add_heading("Sources & Bibliography", level=1)
        all_sources = []
        for r in report.agent_results:
            if r.sources:
                all_sources.extend(r.sources)
        unique_sources = sorted(list(set(all_sources)))
        http_sources = [src for src in unique_sources if src.startswith("http")]
        if not http_sources:
            doc.add_paragraph("No external data source URLs recorded.")
        else:
            for src in http_sources:
                doc.add_paragraph(f"• {src}")

        # Diligence Pipeline Metadata
        doc.add_heading("Diligence Pipeline Metadata", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Agent Name"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Execution Time"
        hdr_cells[3].text = "Error Info"
        for r in report.agent_results:
            row_cells = table.add_row().cells
            row_cells[0].text = r.agent_name
            row_cells[1].text = str(r.status.value if hasattr(r.status, 'value') else r.status)
            row_cells[2].text = f"{r.duration_ms} ms"
            row_cells[3].text = r.error or "None"
                
        doc.save(filepath)
        return filepath

    def render_html(self, report: DiligenceReport) -> str:
        """Generate a beautiful responsive HTML report and return its filepath."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = report.startup_name.lower().replace(' ', '_')
        filepath = os.path.join(self.output_dir, f"{safe_name}_{timestamp}.html")
        
        # HTML Template with beautiful inline CSS (modern layout, gradients, cards,Outfit/Plus Jakarta Font)
        html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>VentureMind Due Diligence: {{ report.startup_name }}</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Plus+Jakarta+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-primary: #0b0f19;
            --bg-secondary: #111827;
            --bg-card: #1f2937;
            --text-primary: #f9fafb;
            --text-secondary: #9ca3af;
            --accent-blue: #3b82f6;
            --accent-purple: #8b5cf6;
            --accent-green: #10b981;
            --accent-amber: #f59e0b;
            --accent-red: #ef4444;
            --border-color: rgba(255, 255, 255, 0.08);
            --card-shadow: 0 10px 30px -10px rgba(0, 0, 0, 0.5);
            --gradient-accent: linear-gradient(135deg, var(--accent-blue) 0%, var(--accent-purple) 100%);
        }
        
        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }
        
        body {
            font-family: 'Plus Jakarta Sans', sans-serif;
            background-color: var(--bg-primary);
            color: var(--text-primary);
            line-height: 1.6;
            padding: 40px 20px;
        }
        
        .container {
            max-width: 1100px;
            margin: 0 auto;
        }
        
        header {
            background: var(--bg-secondary);
            border: 1px solid var(--border-color);
            padding: 40px;
            border-radius: 24px;
            margin-bottom: 30px;
            position: relative;
            overflow: hidden;
            box-shadow: var(--card-shadow);
        }
        
        header::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 6px;
            background: var(--gradient-accent);
        }
        
        .header-meta {
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 20px;
            margin-top: 20px;
            border-top: 1px solid var(--border-color);
            padding-top: 20px;
        }
        
        h1 {
            font-family: 'Outfit', sans-serif;
            font-size: 2.5rem;
            font-weight: 800;
            letter-spacing: -0.03em;
            background: linear-gradient(120deg, #ffffff 30%, #a5b4fc 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 10px;
        }
        
        .subtitle {
            color: var(--text-secondary);
            font-size: 1.1rem;
        }
        
        .badge {
            display: inline-flex;
            align-items: center;
            padding: 10px 20px;
            border-radius: 50px;
            font-weight: 700;
            font-size: 0.95rem;
            letter-spacing: 0.02em;
            text-transform: uppercase;
            box-shadow: 0 4px 12px rgba(0,0,0,0.25);
        }
        
        .badge-exceptional { background: rgba(139, 92, 246, 0.15); color: #c084fc; border: 1px solid rgba(139, 92, 246, 0.3); }
        .badge-strong { background: rgba(16, 185, 129, 0.15); color: #34d399; border: 1px solid rgba(16, 185, 129, 0.3); }
        .badge-moderate { background: rgba(245, 158, 11, 0.15); color: #fbbf24; border: 1px solid rgba(245, 158, 11, 0.3); }
        .badge-risk { background: rgba(239, 68, 68, 0.15); color: #f87171; border: 1px solid rgba(239, 68, 68, 0.3); }
        .badge-not { background: rgba(239, 68, 68, 0.25); color: #f87171; border: 2px solid #ef4444; }
        
        .section-card {
            background: var(--bg-secondary);
            border: 1px solid var(--border-color);
            padding: 35px;
            border-radius: 20px;
            margin-bottom: 30px;
            box-shadow: var(--card-shadow);
        }
        
        h2 {
            font-family: 'Outfit', sans-serif;
            font-size: 1.75rem;
            font-weight: 700;
            margin-bottom: 20px;
            display: flex;
            align-items: center;
            gap: 12px;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 15px;
        }
        
        .grid-3 {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }
        
        .metric-card {
            background: rgba(255, 255, 255, 0.03);
            border: 1px solid var(--border-color);
            padding: 20px;
            border-radius: 12px;
            text-align: center;
        }
        
        .metric-val {
            font-family: 'Outfit', sans-serif;
            font-size: 1.8rem;
            font-weight: 700;
            color: var(--accent-blue);
            margin-top: 5px;
        }
        
        .metric-label {
            color: var(--text-secondary);
            font-size: 0.85rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }
        
        .risk-list {
            list-style: none;
            display: flex;
            flex-direction: column;
            gap: 12px;
        }
        
        .risk-item {
            background: rgba(239, 68, 68, 0.05);
            border-left: 4px solid var(--accent-red);
            padding: 15px 20px;
            border-radius: 0 8px 8px 0;
            display: flex;
            align-items: center;
            gap: 12px;
        }
        
        .risk-item-clean {
            background: rgba(16, 185, 129, 0.05);
            border-left: 4px solid var(--accent-green);
            padding: 15px 20px;
            border-radius: 0 8px 8px 0;
            display: flex;
            align-items: center;
            gap: 12px;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
            background: rgba(255, 255, 255, 0.01);
            border-radius: 8px;
            overflow: hidden;
        }
        
        th, td {
            padding: 14px 18px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        
        th {
            background: rgba(255, 255, 255, 0.04);
            font-weight: 600;
            color: var(--text-primary);
        }
        
        tr:last-child td {
            border-bottom: none;
        }
        
        .timeline {
            display: flex;
            flex-direction: column;
            gap: 15px;
            margin-top: 15px;
        }
        
        .timeline-item {
            display: flex;
            gap: 20px;
            border-left: 2px solid var(--accent-purple);
            padding-left: 20px;
            margin-left: 10px;
            position: relative;
        }
        
        .timeline-item::before {
            content: '';
            position: absolute;
            left: -6px;
            top: 6px;
            width: 10px;
            height: 10px;
            border-radius: 50%;
            background: var(--accent-purple);
        }
        
        .timeline-round {
            font-weight: 700;
            color: var(--text-primary);
        }
        
        .timeline-amount {
            color: var(--accent-green);
            font-weight: 600;
        }
        
        .timeline-date {
            color: var(--text-secondary);
            font-size: 0.9rem;
        }
        
        .badge-status {
            display: inline-block;
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 0.75rem;
            font-weight: 700;
            text-transform: uppercase;
        }
        .status-success { background: rgba(16, 185, 129, 0.2); color: #34d399; }
        .status-failed { background: rgba(239, 68, 68, 0.2); color: #f87171; }
        .status-timeout { background: rgba(245, 158, 11, 0.2); color: #fbbf24; }
        
        .source-link {
            color: var(--accent-blue);
            text-decoration: none;
            display: inline-block;
            margin-bottom: 8px;
        }
        .source-link:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>{{ report.startup_name }} Due Diligence</h1>
            <div class="subtitle">VentureMind Multi-Agent Analysis Report</div>
            <div class="header-meta">
                <div>
                    <strong>Date Generated:</strong> {{ formatted_date }}
                </div>
                <div>
                    <span class="badge {{ badge_class }}">{{ score_label }}</span>
                </div>
            </div>
        </header>

        <!-- Executive Summary -->
        <div class="section-card">
            <h2>📝 Executive Summary</h2>
            <p>{{ report.summary }}</p>
        </div>

        <!-- Risks -->
        <div class="section-card">
            <h2>⚠️ Risk Assessment</h2>
            <ul class="risk-list">
                {% if report.risk_flags %}
                    {% for flag in report.risk_flags %}
                        <li class="risk-item">
                            <span>🚨</span> {{ flag }}
                        </li>
                    {% endfor %}
                {% else %}
                    <li class="risk-item-clean">
                        <span>✅</span> No significant risk flags identified.
                    </li>
                {% endif %}
            </ul>
        </div>

        <!-- Market -->
        <div class="section-card">
            <h2>📊 Market Opportunity</h2>
            {% if report.market %}
                <div class="grid-3">
                    <div class="metric-card">
                        <div class="metric-label">TAM (Total Addressable)</div>
                        <div class="metric-val">${{ "{:,.0f}".format(report.market.tam_usd) if report.market.tam_usd else "N/A" }}</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-label">SAM (Serviceable Addressable)</div>
                        <div class="metric-val">${{ "{:,.0f}".format(report.market.sam_usd) if report.market.sam_usd else "N/A" }}</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-label">SOM (Serviceable Obtainable)</div>
                        <div class="metric-val">${{ "{:,.0f}".format(report.market.som_usd) if report.market.som_usd else "N/A" }}</div>
                    </div>
                </div>
                <div style="margin-top: 25px;">
                    <strong>Compound Annual Growth Rate (CAGR):</strong> {{ report.market.cagr_pct }}%
                </div>
                <div style="margin-top: 15px;">
                    <strong>Key Industry Trends:</strong>
                    <ul style="margin-left: 20px; margin-top: 5px;">
                        {% for trend in report.market.key_trends %}
                            <li>{{ trend }}</li>
                        {% endfor %}
                    </ul>
                </div>
            {% else %}
                <p>Market data is unavailable.</p>
            {% endif %}
        </div>

        <!-- Competitors -->
        <div class="section-card">
            <h2>⚔️ Competitive Landscape</h2>
            {% if report.competitors %}
                <div style="margin-bottom: 20px;">
                    <strong>Differentiation Score:</strong> {{ differentiation_pct }}%
                </div>
                <div style="margin-bottom: 20px;">
                    <strong>Positioning Summary:</strong><br>
                    {{ report.competitors.positioning_summary }}
                </div>
                <strong>Direct Competitor Matrix:</strong>
                <table>
                    <thead>
                        <tr>
                            <th>Competitor Name</th>
                            <th>Founded Year</th>
                            <th>Total Funding</th>
                            <th>Market Share</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for c in report.competitors.competitors %}
                            <tr>
                                <td><strong>{{ c.name }}</strong></td>
                                <td>{{ c.founded_year }}</td>
                                <td style="color: #34d399;">${{ "{:,.0f}".format(c.funding_usd) if c.funding_usd else "unknown" }}</td>
                                <td>{{ "{:.1f}%".format(c.market_share_pct) if c.market_share_pct else "unknown" }}</td>
                            </tr>
                        {% endfor %}
                    </tbody>
                </table>
            {% else %}
                <p>Competitive data is unavailable.</p>
            {% endif %}
        </div>

        <!-- Financials -->
        <div class="section-card">
            <h2>💰 Financial Profile</h2>
            {% if report.financials %}
                <div style="margin-bottom: 20px;">
                    <strong>Estimated Burn Rate:</strong> {{ report.financials.burn_rate_estimate }} | 
                    <strong>Estimated Runway:</strong> {{ report.financials.runway_estimate }}
                </div>
                
                <h3 style="margin-top: 20px; margin-bottom: 10px; font-size: 1.15rem; color: var(--text-secondary);">Financial Signals & Metrics</h3>
                <table>
                    <thead>
                        <tr>
                            <th>Metric</th>
                            <th>Value</th>
                            <th>Period</th>
                            <th>Source</th>
                            <th>Confidence</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for s in report.financials.signals %}
                            <tr>
                                <td><strong>{{ s.metric_name }}</strong></td>
                                <td>{{ s.value }}</td>
                                <td>{{ s.period }}</td>
                                <td>{{ s.source }}</td>
                                <td>{{ "{:.2f}".format(s.confidence) }}</td>
                            </tr>
                        {% endfor %}
                    </tbody>
                </table>

                <h3 style="margin-top: 25px; margin-bottom: 15px; font-size: 1.15rem; color: var(--text-secondary);">Funding Timeline</h3>
                <div class="timeline">
                    {% for r in report.financials.funding_rounds %}
                        <div class="timeline-item">
                            <div>
                                <span class="timeline-round">{{ r.get('round', 'Round') }}</span> - 
                                <span class="timeline-amount">{{ r.get('amount_str', 'unknown') }}</span>
                                <div class="timeline-date">{{ r.get('date_str', 'unknown') }} (via {{ r.get('source', 'unknown') }})</div>
                            </div>
                        </div>
                    {% endfor %}
                </div>
            {% else %}
                <p>Financial data is unavailable.</p>
            {% endif %}
        </div>

        <!-- Legal -->
        <div class="section-card">
            <h2>⚖️ Legal Standing</h2>
            {% if report.legal %}
                <div style="margin-bottom: 20px;">
                    <strong>Incorporation Status:</strong> {{ report.legal.incorporation_status }} | 
                    <strong>Patent Count:</strong> {{ report.legal.patent_count }} | 
                    <strong>Trademark Count:</strong> {{ report.legal.trademark_count }}
                </div>
                
                <h3 style="margin-top: 20px; margin-bottom: 10px; font-size: 1.15rem; color: var(--text-secondary);">Legal & Compliance Flags</h3>
                <table>
                    <thead>
                        <tr>
                            <th>Issue / Flag Type</th>
                            <th>Severity</th>
                            <th>Description</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for f in report.legal.flags %}
                            <tr>
                                <td><strong>{{ f.flag_type }}</strong></td>
                                <td>
                                    <span class="badge-status {{ 'status-failed' if f.severity.upper() == 'HIGH' else 'status-timeout' if f.severity.upper() == 'MEDIUM' else 'status-success' }}">
                                        {{ f.severity }}
                                    </span>
                                </td>
                                <td>{{ f.description }}</td>
                            </tr>
                        {% else %}
                            <tr>
                                <td colspan="3">No active legal or regulatory flags identified.</td>
                            </tr>
                        {% endfor %}
                    </tbody>
                </table>
            {% else %}
                <p>Legal standing data is unavailable.</p>
            {% endif %}
        </div>

        <!-- Sources -->
        <div class="section-card">
            <h2>🔗 Sources & Bibliography</h2>
            {% set all_sources = [] %}
            {% for r in report.agent_results %}
                {% if r.sources %}
                    {% for src in r.sources %}
                        {% if src.startswith("http") %}
                            {% if src not in all_sources %}
                                {% set _ = all_sources.append(src) %}
                            {% endif %}
                        {% endif %}
                    {% endfor %}
                {% endif %}
            {% endfor %}
            
            {% if all_sources %}
                <div style="display: flex; flex-direction: column;">
                    {% for src in all_sources|sort %}
                        <a href="{{ src }}" class="source-link" target="_blank">{{ src }}</a>
                    {% endfor %}
                </div>
            {% else %}
                <p>No external data source URLs recorded.</p>
            {% endif %}
        </div>

        <!-- Pipeline Metadata -->
        <div class="section-card">
            <h2>⚙️ Diligence Pipeline Metadata</h2>
            <table>
                <thead>
                    <tr>
                        <th>Agent Name</th>
                        <th>Status</th>
                        <th>Execution Time</th>
                        <th>Error Info</th>
                    </tr>
                </thead>
                <tbody>
                    {% for r in report.agent_results %}
                        <tr>
                            <td><strong>{{ r.agent_name }}</strong></td>
                            <td>
                                <span class="badge-status {{ 'status-success' if r.status.value == 'SUCCESS' or r.status == 'SUCCESS' else 'status-failed' if r.status.value == 'FAILED' or r.status == 'FAILED' else 'status-timeout' }}">
                                    {{ r.status.value if hasattr(r.status, 'value') else r.status }}
                                </span>
                            </td>
                            <td>{{ r.duration_ms }} ms</td>
                            <td>{{ r.error or 'None' }}</td>
                        </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    </div>
</body>
</html>
"""
        
        # Render helper variables
        score = report.investment_score
        if 0.0 <= score < 3.0:
            badge_class = "badge-not"
            score_label = f"{score:.1f}/10 — NOT INVESTABLE"
        elif 3.0 <= score < 5.0:
            badge_class = "badge-risk"
            score_label = f"{score:.1f}/10 — HIGH RISK"
        elif 5.0 <= score < 7.0:
            badge_class = "badge-moderate"
            score_label = f"{score:.1f}/10 — MODERATE OPPORTUNITY"
        elif 7.0 <= score < 9.0:
            badge_class = "badge-strong"
            score_label = f"{score:.1f}/10 — STRONG OPPORTUNITY"
        else:
            badge_class = "badge-exceptional"
            score_label = f"{score:.1f}/10 — EXCEPTIONAL OPPORTUNITY"

        diff_pct = 0
        if report.competitors:
            diff_pct = int(report.competitors.differentiation_score * 100)

        formatted_date = datetime.now().strftime("%B %d, %Y")

        if HAS_JINJA:
            template = Template(html_template)
            rendered = template.render(
                report=report,
                formatted_date=formatted_date,
                badge_class=badge_class,
                score_label=score_label,
                differentiation_pct=diff_pct,
                hasattr=hasattr
            )
        else:
            # Fallback simple render if Jinja2 is missing (though we will make sure it's installed)
            rendered = f"<html><body><h1>{report.startup_name} Due Diligence Report</h1><p>Score: {score:.1f}/10</p><p>{report.summary}</p></body></html>"

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(rendered)
        return filepath

    def render_executive_summary(self, report: DiligenceReport) -> str:
        """Short 1-page Markdown summary."""
        badge = self._render_score_badge(report.investment_score)
        top_risks = "\n".join(f"- ⚠️ {r}" for r in report.risk_flags[:3]) if report.risk_flags else "No major risks identified."
        
        return (
            f"# Executive Summary: {report.startup_name}\n\n"
            f"**Investment Recommendation:** {badge}\n\n"
            f"{report.summary}\n\n"
            "#### Top Risk Flags\n"
            f"{top_risks}\n"
        )

    def _render_market_section(self, market: MarketData | None) -> str:
        """Render the market research and TAM/SAM/SOM findings section."""
        if not market:
            return "### Market Opportunity\nData unavailable.\n"
        
        table = (
            "| Metric | Value (USD) |\n"
            "| --- | --- |\n"
            f"| TAM (Total Addressable Market) | ${market.tam_usd:,.2f} |\n"
            f"| SAM (Serviceable Addressable Market) | ${market.sam_usd:,.2f} |\n"
            f"| SOM (Serviceable Obtainable Market) | ${market.som_usd:,.2f} |\n"
        )
        
        trends_list = "\n".join(f"- {trend}" for trend in market.key_trends)
        
        return (
            "### Market Opportunity\n\n"
            f"**Compound Annual Growth Rate (CAGR):** {market.cagr_pct:.1f}%\n\n"
            "#### Market Sizing\n"
            f"{table}\n\n"
            "#### Key Industry Trends\n"
            f"{trends_list}\n"
        )

    def _render_competitor_section(self, competitors: CompetitorLandscape | None) -> str:
        """Render the competitive landscape and positioning section."""
        if not competitors:
            return "### Competitive Landscape\nData unavailable.\n"
        
        comp_rows = []
        for c in competitors.competitors:
            funding = f"${c.funding_usd:,.2f}" if c.funding_usd else "unknown"
            share = f"{c.market_share_pct:.1f}%" if c.market_share_pct else "unknown"
            comp_rows.append(f"| {c.name} | {c.founded_year} | {funding} | {share} |")
            
        table = (
            "| Competitor Name | Founded Year | Total Funding | Market Share |\n"
            "| --- | --- | --- | --- |\n" +
            "\n".join(comp_rows)
        )
        
        differentiation_pct = int(competitors.differentiation_score * 100)
        
        return (
            "### Competitive Landscape\n\n"
            f"**Differentiation Score:** {differentiation_pct}%\n\n"
            f"**Positioning Summary:**\n{competitors.positioning_summary}\n\n"
            "#### Direct Competitor Matrix\n"
            f"{table}\n"
        )

    def _render_financial_section(self, financials: FinancialProfile | None) -> str:
        """Render the financial standing and metrics overview section."""
        if not financials:
            return "### Financial Profile\nData unavailable.\n"
        
        round_rows = []
        for r in financials.funding_rounds:
            round_rows.append(f"- **{r.get('round', 'Unknown')}**: {r.get('amount_str', 'unknown')} on {r.get('date_str', 'unknown')} (Source: {r.get('source', 'unknown')})")
            
        rounds_timeline = "\n".join(round_rows) if round_rows else "No funding rounds recorded."
        
        signals_rows = []
        for s in financials.signals:
            signals_rows.append(f"| {s.metric_name} | {s.value} | {s.period} | {s.source} | {s.confidence:.2f} |")
            
        signals_table = (
            "| Metric | Value | Period | Source | Confidence |\n"
            "| --- | --- | --- | --- | --- |\n" +
            "\n".join(signals_rows)
        )
        
        return (
            "### Financial Profile\n\n"
            f"**Estimated Burn Rate:** {financials.burn_rate_estimate}\n"
            f"**Estimated Runway:** {financials.runway_estimate}\n\n"
            "#### Financial Signals & Metrics\n"
            f"{signals_table}\n\n"
            "#### Funding Rounds Timeline\n"
            f"{rounds_timeline}\n"
        )

    def _render_legal_section(self, legal: LegalProfile | None) -> str:
        """Render the incorporation, IP, litigation, and regulatory flags section."""
        if not legal:
            return "### Legal Standing\nData unavailable.\n"
        
        flag_rows = []
        for f in legal.flags:
            flag_rows.append(f"| {f.flag_type} | {f.severity} | {f.description} |")
            
        flags_table = (
            "| Issue / Flag Type | Severity | Description |\n"
            "| --- | --- | --- |\n" +
            "\n".join(flag_rows)
        ) if flag_rows else "No active legal or regulatory flags identified."
        
        return (
            "### Legal Standing\n\n"
            f"**Incorporation Status:** {legal.incorporation_status}\n"
            f"**Patent Count Estimate:** {legal.patent_count}\n"
            f"**Trademark Count Estimate:** {legal.trademark_count}\n\n"
            "#### Legal & Compliance Flags\n"
            f"{flags_table}\n"
        )

    def _render_score_badge(self, score: float) -> str:
        """Returns visual score badge string like '⭐ 7.2/10 — STRONG OPPORTUNITY'"""
        if 0.0 <= score < 3.0:
            return f"🔴 {score:.1f}/10 — NOT INVESTABLE"
        elif 3.0 <= score < 5.0:
            return f"🟡 {score:.1f}/10 — HIGH RISK"
        elif 5.0 <= score < 7.0:
            return f"🟠 {score:.1f}/10 — MODERATE OPPORTUNITY"
        elif 7.0 <= score < 9.0:
            return f"🟢 {score:.1f}/10 — STRONG OPPORTUNITY"
        else:
            return f"⭐ {score:.1f}/10 — EXCEPTIONAL OPPORTUNITY"

    def _format_risk_flags(self, flags: list[str]) -> str:
        """Format a list of risk flags into a structured Markdown block."""
        if not flags:
            return "No significant risk flags identified."
        return "\n".join(f"- ⚠️ {flag}" for flag in flags)
